import os
import datetime
import time
import io
import docx
import openai
import boto3
from dotenv import load_dotenv
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from duckduckgo_search import DDGS
from urllib.parse import urlparse
load_dotenv()

# Set up Boto3 S3 client
s3_client = boto3.client(
    "s3",
    aws_access_key_id=os.getenv(AWS_ACCESS_KEY_ID),
    aws_secret_access_key=os.getenv(AWS_SECRET_ACCESS_KEY),
    region_name="ap-south-1" 
)

# Set up Selenium WebDriver (Headless Mode)
def get_driver():
    options = Options()
    options.add_argument("--headless")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--ignore-certificate-errors")
    options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    return driver

# Function to scrape page content using Selenium
def scrape_url(url):
    driver = get_driver()
    try:
        driver.get(url)
        time.sleep(5)
        page_text = driver.find_element(By.TAG_NAME, "body").text
        return page_text
    except Exception as e:
        print(f"Error scraping {url}: {e}")
        return ""
    finally:
        driver.quit()

# Function to extract detailed relevant content using OpenAI
def extract_relevant_content(keyword, text):
    
    openai.api_key = OPENAI_API_KEY 

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Extract the most relevant, detailed, and structured information from the given text. Ensure the extracted content is at least 2 pages long."},
            {"role": "user", "content": f"""Extract all the **important details** related to '{keyword}' from the following text.
                
                - Summarize key points **clearly**.
                - Use **headings and subheadings** for structure.
                - Ensure the **content is at least 1-2 pages** long.
                - Expand on important **facts, statistics, and examples**.
                - If relevant, provide **historical context, latest trends, and expert opinions**.

                TEXT:
                {text}
            """}
        ],
        max_tokens=3000 
    )

    return response["choices"][0]["message"]["content"].strip()

# Function to save extracted content as DOCX and upload to S3
def save_to_s3(url, text, date):
   
    domain = urlparse(url).netloc.replace("www.", "").replace(".", "_")
    file_name = f"{domain}_{date}.docx"
    
    # Create a DOCX file in memory
    doc_buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_heading(f"Extracted Content from {url}", level=1)
    doc.add_paragraph(f"Date Extracted: {date}\n\n")
    doc.add_paragraph(text)
    doc.save(doc_buffer)
    
    # Move cursor to the start of the buffer before uploading
    doc_buffer.seek(0)

    # Upload to S3
    s3_key = f"scraped_data/{file_name}"
    s3_client.put_object(
        Bucket=S3_BUCKET_NAME,
        Key=s3_key,
        Body=doc_buffer.getvalue(),
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    print(f"Uploaded {file_name} to S3: s3://{S3_BUCKET_NAME}/{s3_key}")

# Function to perform search, filter trusted sources, scrape, and save results
def search_and_scrape(keywords, trusted_sites):
    today_date = datetime.date.today().strftime("%d-%m-%Y")

    for keyword in keywords:
        print(f" Searching for: {keyword}")

        query = f"{keyword} {today_date}"
        search_results = DDGS().text(query, max_results=10)
        
        print("\n Found URLs:")
        for result in search_results:
            print(result["href"])

        # Filter results based on trusted sources
        filtered_urls = [result["href"] for result in search_results if any(site in urlparse(result["href"]).netloc for site in trusted_sites)]
        
        if not filtered_urls:
            print(f"No trusted sources found for '{keyword}'\n")
            continue
        
        print(f"{len(filtered_urls)} trusted results for '{keyword}'\n")

        saved = False  # Track if any file is saved

        for url in filtered_urls:
            extracted_text = scrape_url(url)
            if extracted_text:
                relevant_content = extract_relevant_content(keyword, extracted_text)
                save_to_s3(url, relevant_content, today_date)
                saved = True
        
        if not saved:
            print(f"No relevant content found for '{keyword}'\n")



keywords = ["Risk Management", "Cyber Crimes", "Inflation", "Geo-political tension"]
trusted_sites = ["bloomberg.com", "forbes.com", "reuters.com", "diplomatist.com"]


search_and_scrape(keywords, trusted_sites)
